from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from datetime import date, timedelta, datetime
import calendar
import pickle as p
from io import BytesIO
import requests
from icalendar import Calendar

year = date.today().year

def fetch_holidays(pYear, url):
    r = requests.get(url)
    r.raise_for_status()

    cal = Calendar.from_ical(r.text)
    holidays_dict = {}

    for event in cal.walk('VEVENT'):
        dt = event.get('dtstart').dt
        if isinstance(dt, datetime):
            dt = dt.date()
        if dt.year == pYear:
            holidays_dict[dt] = str(event.get('summary'))

    return holidays_dict

def create_month_doc(months=None, showHolidays=True, url="https://calendar.google.com/calendar/ical/de.german%23holiday@group.v.calendar.google.com/public/basic.ics"):
    if months is None:
        months = [date.today().month]
    if isinstance(months, int):
        months = [months]

    weekend_color = "92D050"
    de_holidays = fetch_holidays(year, url)
    doc = Document()

    for i, month in enumerate(months):
        month_name = calendar.month_name[month]
        num_days = calendar.monthrange(year, month)[1]

        doc.add_heading(f"{month_name} {year}", level=1)

        #create table days
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Datum"
        hdr_cells[1].text = "Lorem Ipsum"

        for day in range(1, num_days +1):
            current_date = date(year, month, day)
            row_cells = table.add_row().cells

            if current_date in de_holidays and showHolidays:
                row_cells[0].text = current_date.strftime("%d.") + " " + str(de_holidays[current_date])
            else:
                row_cells[0].text = current_date.strftime("%d.")

            row_cells[1].text = ''

            # fill weekends (sa = 5; su = 6)
            if current_date.weekday() >= 5:
                tc = row_cells[0]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), weekend_color)
                tcPr.append(shd)
            
        for cell in table.columns[0].cells:
            cell.width = Cm(1.0)
        for cell in table.columns[1].cells:
            cell.width = Cm((7.74)*2-1.0)

        if i < len(months) -1:
            doc.add_page_break()

    #save to mem
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    if len(months) == 1:
        filename = f"{calendar.month_name[months[0]]}_{year}.docx"
    else:
        filename = f"Monate_{calendar.month_name[months[0]]}-{calendar.month_name[months[-1]]}_{year}.docx"

    return file_stream, filename

def create_month_html(months=None, showHolidays=True, url="https://calendar.google.com/calendar/ical/de.german%23holiday@group.v.calendar.google.com/public/basic.ics"):
    if months is None: 
        months = [date.today().month]
    if isinstance(months, int):
        months = [months]

    de_holidays = fetch_holidays(year, url)

    html = """
    <style>
        table { border-collapse: collapse; width: 100%; margin-bottom: 30px; }
        th, td { border: 1px solid black; padding: 4px; }
        .weekend { background-color: #92D050; }
    </style>
    """

    for month in months:
        month_name = calendar.month_name[month]
        num_days = calendar.monthrange(year, month)[1]

        html += f"<h3 class='h_month_name'>{month_name} {year}</h3>"
        html += """
        <table>
            <tr>
                <th style="width: 50px;">Datum</th>
                <th>Lorem Ipsum</th>
            </tr>
        """

        for day in range(1, num_days +1):
            current_date = date(year, month, day)
            weekend = "weekend" if current_date.weekday() >= 5 else ""

            holiday = str(de_holidays[current_date]) if current_date in de_holidays and showHolidays == True else ""

            html += f"""
            <tr>
                <td class="{weekend}">{current_date.strftime('%d.') + holiday}</td>
                <td></td>
            </tr>
            """

        html += "</table>"
    return html